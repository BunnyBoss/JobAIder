from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Annotated
from io import BytesIO

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

from app.core.config import get_settings
from app.db import decode_json, encode_json, get_db, init_db, now_iso
from app.modules.document_ingestion import normalize_uploaded_name, read_document
from app.modules.gap_analyzer import analyze_gap
from app.modules.improvement_planner import generate_improvement_plan
from app.modules.interview_coach import evaluate_answer, first_question, generate_session_feedback
from app.modules.obsidian_parser import scan_vault
from app.modules.profile_builder import build_profile
from app.modules.resume_generator import generate_generic_resumes, generate_tailored_resume
from app.modules.role_analyzer import analyze_role
from app.schemas import (
    BuildProfileRequest,
    GapAnalysisUpdateRequest,
    GapAnalysisRequest,
    InterviewAnswerRequest,
    InterviewStartRequest,
    ObsidianScanRequest,
    ProfileUpdateRequest,
    ProviderSettings,
    ResumeGenerateRequest,
    ResumeRequest,
    ResumeUpdateRequest,
    RoleAnalysisUpdateRequest,
    RoleAnalysisRequest,
)

app = FastAPI(title="JobRadar API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def startup() -> None:
    init_db()


def _row_json(table: str, row_id: int, json_column: str) -> dict:
    with get_db() as conn:
        row = conn.execute(f"SELECT {json_column} FROM {table} WHERE id = ?", (row_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail=f"{table} row not found")
    return decode_json(row[json_column])


def _insert_document(source_type: str, name: str, content: str, path: str | None = None) -> int:
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO documents (source_type, name, path, content, metadata_json, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (source_type, name, path, content, "{}", now_iso()),
        )
        return int(cur.lastrowid)


def _resume_as_profile(resume_id: int) -> tuple[int, dict]:
    with get_db() as conn:
        row = conn.execute(
            "SELECT profile_id, role_analysis_id, title, kind, markdown, metadata_json FROM resumes WHERE id = ?",
            (resume_id,),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Resume not found")
    return int(row["profile_id"]), {
        "summary": f"{row['title']} ({row['kind']})",
        "resume_id": resume_id,
        "resume_kind": row["kind"],
        "resume_title": row["title"],
        "resume_markdown": row["markdown"],
        "source_role_analysis_id": row["role_analysis_id"],
        "metadata": decode_json(row["metadata_json"]),
    }


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _markdown_to_minimal_pdf(markdown: str) -> bytes:
    lines = markdown.replace("\r\n", "\n").splitlines()
    content = ["BT", "/F1 10 Tf", "50 760 Td", "14 TL"]
    for line in lines[:52]:
        content.append(f"({_pdf_escape(line[:95])}) Tj")
        content.append("T*")
    content.append("ET")
    stream = "\n".join(content).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    pdf = BytesIO()
    pdf.write(b"%PDF-1.4\n")
    offsets = [0]
    for idx, obj in enumerate(objects, start=1):
        offsets.append(pdf.tell())
        pdf.write(f"{idx} 0 obj\n".encode())
        pdf.write(obj)
        pdf.write(b"\nendobj\n")
    xref = pdf.tell()
    pdf.write(f"xref\n0 {len(objects) + 1}\n".encode())
    pdf.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.write(f"{offset:010d} 00000 n \n".encode())
    pdf.write(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return pdf.getvalue()


@app.get("/api/health")
def health() -> dict:
    return {"status": "ok"}


@app.get("/api/settings")
def read_settings() -> dict:
    settings = get_settings()
    with get_db() as conn:
        rows = conn.execute(
            "SELECT key, value FROM settings WHERE key IN ('model_name', 'openai_api_key', 'openai_base_url')"
        ).fetchall()
    stored = {row["key"]: row["value"] for row in rows}
    return {
        "model_name": stored.get("model_name") or settings.model_name,
        "openai_base_url": stored.get("openai_base_url") or settings.openai_base_url,
        "openai_api_key": stored.get("openai_api_key") or settings.openai_api_key,
        "db_path": str(settings.resolved_db_path),
    }


@app.put("/api/settings")
def update_settings(payload: ProviderSettings) -> dict:
    values = payload.model_dump()
    with get_db() as conn:
        for key, value in values.items():
            conn.execute(
                """
                INSERT INTO settings (key, value, updated_at)
                VALUES (?, ?, ?)
                ON CONFLICT(key) DO UPDATE SET value = excluded.value, updated_at = excluded.updated_at
                """,
                (key, value, now_iso()),
            )
    return {"saved": True, "settings": values}


@app.post("/api/profile/obsidian")
def ingest_obsidian(payload: ObsidianScanRequest) -> dict:
    try:
        documents = scan_vault(payload.vault_path)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    saved = []
    for doc in documents:
        doc_id = _insert_document(doc["source_type"], doc["name"], doc["content"], doc["path"])
        saved.append(
            {
                "id": doc_id,
                "name": doc["name"],
                "path": doc["path"],
                "source_type": doc["source_type"],
                "characters": len(doc["content"]),
            }
        )
    return {"documents": saved}


@app.post("/api/profile/files")
async def upload_profile_files(files: Annotated[list[UploadFile], File(...)]) -> dict:
    saved = []
    for upload in files:
        name = normalize_uploaded_name(upload.filename or "upload")
        suffix = Path(name).suffix
        with NamedTemporaryFile(suffix=suffix, delete=True) as temp:
            temp.write(await upload.read())
            temp.flush()
            content = read_document(Path(temp.name))
        doc_id = _insert_document("upload", name, content)
        saved.append({"id": doc_id, "name": name, "source_type": "upload", "characters": len(content)})
    return {"documents": saved}


@app.post("/api/profile/build")
async def build_master_profile(payload: BuildProfileRequest) -> dict:
    parts = [payload.raw_text] if payload.raw_text else []
    if payload.document_ids:
        placeholders = ",".join("?" for _ in payload.document_ids)
        with get_db() as conn:
            rows = conn.execute(
                f"SELECT name, content FROM documents WHERE id IN ({placeholders})",
                payload.document_ids,
            ).fetchall()
        parts.extend(f"# {row['name']}\n{row['content']}" for row in rows)

    if not parts:
        raise HTTPException(status_code=400, detail="Provide document_ids or raw_text")

    profile = await build_profile("\n\n".join(parts))
    timestamp = now_iso()
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO profiles (name, summary, profile_json, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            ("Master Profile", profile.get("summary", ""), encode_json(profile), timestamp, timestamp),
        )
        profile_id = int(cur.lastrowid)
    return {"id": profile_id, "profile": profile}


@app.get("/api/profiles/latest")
def latest_profile() -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM profiles ORDER BY updated_at DESC LIMIT 1").fetchone()
    if not row:
        return {"profile": None}
    return {"id": row["id"], "profile": decode_json(row["profile_json"])}


@app.get("/api/profiles")
def list_profiles() -> dict:
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT id, name, summary, profile_json, created_at, updated_at
            FROM profiles ORDER BY updated_at DESC
            """
        ).fetchall()
    return {
        "profiles": [
            {
                "id": row["id"],
                "name": row["name"],
                "summary": row["summary"],
                "profile": decode_json(row["profile_json"]),
                "created_at": row["created_at"],
                "updated_at": row["updated_at"],
            }
            for row in rows
        ]
    }


@app.get("/api/profile/{profile_id}")
def get_profile(profile_id: int) -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM profiles WHERE id = ?", (profile_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Profile not found")
    return {
        "id": row["id"],
        "name": row["name"],
        "summary": row["summary"],
        "profile": decode_json(row["profile_json"]),
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }


@app.put("/api/profile/{profile_id}")
def update_profile(profile_id: int, payload: ProfileUpdateRequest) -> dict:
    summary = str(payload.profile.get("summary", ""))
    with get_db() as conn:
        row = conn.execute("SELECT id FROM profiles WHERE id = ?", (profile_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Profile not found")
        conn.execute(
            """
            UPDATE profiles
            SET name = COALESCE(?, name), summary = ?, profile_json = ?, updated_at = ?
            WHERE id = ?
            """,
            (payload.name, summary, encode_json(payload.profile), now_iso(), profile_id),
        )
    return {"saved": True, "id": profile_id}


@app.delete("/api/profile/{profile_id}")
def delete_profile(profile_id: int) -> dict:
    with get_db() as conn:
        cur = conn.execute("DELETE FROM profiles WHERE id = ?", (profile_id,))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Profile not found")
    return {"deleted": True, "id": profile_id}


@app.post("/api/role/analyze")
async def create_role_analysis(payload: RoleAnalysisRequest) -> dict:
    from app.modules.role_analyzer import extract_role_metadata
    
    # Extract title and company from content if not provided or are defaults
    if payload.title == "Target Role" or payload.company == "Unknown Company":
        metadata = await extract_role_metadata(payload.content)
        title = metadata["title"] if payload.title == "Target Role" else payload.title
        company = metadata["company"] if payload.company == "Unknown Company" else payload.company
    else:
        title = payload.title
        company = payload.company
    
    analysis = await analyze_role(payload.content)
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO role_analyses (title, company, source_text, analysis_json, created_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (title, company, payload.content, encode_json(analysis), now_iso()),
        )
        role_id = int(cur.lastrowid)
    return {"id": role_id, "title": title, "company": company, "analysis": analysis}


@app.get("/api/roles")
def list_role_analyses() -> dict:
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT id, title, company, source_text, analysis_json, created_at
            FROM role_analyses ORDER BY created_at DESC
            """
        ).fetchall()
    return {
        "roles": [
            {
                "id": row["id"],
                "title": row["title"],
                "company": row["company"],
                "source_text": row["source_text"],
                "analysis": decode_json(row["analysis_json"]),
                "created_at": row["created_at"],
            }
            for row in rows
        ]
    }


@app.get("/api/role/{role_analysis_id}")
def get_role_analysis(role_analysis_id: int) -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM role_analyses WHERE id = ?", (role_analysis_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Role analysis not found")
    return {
        "id": row["id"],
        "title": row["title"],
        "company": row["company"],
        "source_text": row["source_text"],
        "analysis": decode_json(row["analysis_json"]),
        "created_at": row["created_at"],
    }


@app.put("/api/role/{role_analysis_id}")
def update_role_analysis(role_analysis_id: int, payload: RoleAnalysisUpdateRequest) -> dict:
    with get_db() as conn:
        row = conn.execute(
            "SELECT source_text FROM role_analyses WHERE id = ?", (role_analysis_id,)
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Role analysis not found")
        conn.execute(
            """
            UPDATE role_analyses
            SET title = ?, company = ?, source_text = ?, analysis_json = ?
            WHERE id = ?
            """,
            (
                payload.title,
                payload.company,
                payload.source_text if payload.source_text is not None else row["source_text"],
                encode_json(payload.analysis),
                role_analysis_id,
            ),
        )
    return {"saved": True, "id": role_analysis_id}


@app.delete("/api/role/{role_analysis_id}")
def delete_role_analysis(role_analysis_id: int) -> dict:
    with get_db() as conn:
        cur = conn.execute("DELETE FROM role_analyses WHERE id = ?", (role_analysis_id,))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Role analysis not found")
    return {"deleted": True, "id": role_analysis_id}


@app.post("/api/resume/generic")
async def create_generic_resumes(payload: ResumeRequest) -> dict:
    profile = _row_json("profiles", payload.profile_id, "profile_json")
    ats, human = await generate_generic_resumes(profile)
    responses = []
    with get_db() as conn:
        for kind, markdown in [("ats", ats), ("human", human)]:
            # Ensure markdown is a string (handle cases where it might be a dict)
            if isinstance(markdown, dict):
                import json
                markdown = json.dumps(markdown, ensure_ascii=False, indent=2)
            markdown = str(markdown) if markdown is not None else ""
            
            cur = conn.execute(
                """
                INSERT INTO resumes (profile_id, role_analysis_id, kind, title, markdown, metadata_json, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (payload.profile_id, None, kind, f"{kind.upper()} Resume", markdown, "{}", now_iso()),
            )
            responses.append({"id": int(cur.lastrowid), "kind": kind, "markdown": markdown})
    return {"resumes": responses}


@app.post("/api/resume/generate")
async def generate_resume(payload: ResumeGenerateRequest) -> dict:
    profile = _row_json("profiles", payload.profile_id, "profile_json")
    metadata = {}
    if payload.kind == "tailored":
        if payload.role_analysis_id is None:
            raise HTTPException(status_code=400, detail="role_analysis_id is required for tailored")
        role = _row_json("role_analyses", payload.role_analysis_id, "analysis_json")
        result = await generate_tailored_resume(profile, role)
        markdown = result["markdown"]
        metadata = {k: v for k, v in result.items() if k != "markdown"}
        title = "Tailored Resume"
    else:
        ats, human = await generate_generic_resumes(profile)
        markdown = ats if payload.kind == "ats" else human
        title = "ATS Resume" if payload.kind == "ats" else "Human-Friendly Resume"

    # Ensure markdown is a string (handle cases where it might be a dict)
    if isinstance(markdown, dict):
        import json
        markdown = json.dumps(markdown, ensure_ascii=False, indent=2)
    markdown = str(markdown) if markdown is not None else ""

    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO resumes (profile_id, role_analysis_id, kind, title, markdown, metadata_json, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.profile_id,
                payload.role_analysis_id,
                payload.kind,
                title,
                markdown,
                encode_json(metadata),
                now_iso(),
            ),
        )
    return {
        "id": int(cur.lastrowid),
        "profile_id": payload.profile_id,
        "role_analysis_id": payload.role_analysis_id,
        "kind": payload.kind,
        "title": title,
        "markdown": markdown,
        "metadata": metadata,
    }


@app.post("/api/resume/tailored")
async def create_tailored_resume(payload: ResumeRequest) -> dict:
    if payload.role_analysis_id is None:
        raise HTTPException(status_code=400, detail="role_analysis_id is required")
    profile = _row_json("profiles", payload.profile_id, "profile_json")
    role = _row_json("role_analyses", payload.role_analysis_id, "analysis_json")
    result = await generate_tailored_resume(profile, role)
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO resumes (profile_id, role_analysis_id, kind, title, markdown, metadata_json, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload.profile_id,
                payload.role_analysis_id,
                "tailored",
                "Tailored Resume",
                result["markdown"],
                encode_json({k: v for k, v in result.items() if k != "markdown"}),
                now_iso(),
            ),
        )
    return {"id": int(cur.lastrowid), "kind": "tailored", **result}


@app.get("/api/resumes")
def list_resumes() -> dict:
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT id, profile_id, role_analysis_id, kind, title, markdown, metadata_json, created_at
            FROM resumes ORDER BY created_at DESC
            """
        ).fetchall()
    return {
        "resumes": [
            {
                "id": row["id"],
                "profile_id": row["profile_id"],
                "role_analysis_id": row["role_analysis_id"],
                "kind": row["kind"],
                "title": row["title"],
                "markdown": row["markdown"],
                "metadata": decode_json(row["metadata_json"]),
                "created_at": row["created_at"],
            }
            for row in rows
        ]
    }


@app.get("/api/resume/{resume_id}")
def get_resume(resume_id: int) -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM resumes WHERE id = ?", (resume_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {
        "id": row["id"],
        "profile_id": row["profile_id"],
        "role_analysis_id": row["role_analysis_id"],
        "kind": row["kind"],
        "title": row["title"],
        "markdown": row["markdown"],
        "metadata": decode_json(row["metadata_json"]),
        "created_at": row["created_at"],
    }


@app.put("/api/resume/{resume_id}")
def update_resume(resume_id: int, payload: ResumeUpdateRequest) -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT id FROM resumes WHERE id = ?", (resume_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Resume not found")
        conn.execute(
            "UPDATE resumes SET title = COALESCE(?, title), markdown = ? WHERE id = ?",
            (payload.title, payload.markdown, resume_id),
        )
    return {"saved": True, "id": resume_id}


@app.delete("/api/resume/{resume_id}")
def delete_resume(resume_id: int) -> dict:
    with get_db() as conn:
        cur = conn.execute("DELETE FROM resumes WHERE id = ?", (resume_id,))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Resume not found")
    return {"deleted": True, "id": resume_id}


@app.get("/api/resume/{resume_id}/export")
def export_resume(resume_id: int, format: str = "markdown") -> StreamingResponse:
    with get_db() as conn:
        row = conn.execute("SELECT title, markdown FROM resumes WHERE id = ?", (resume_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Resume not found")

    title = row["title"].lower().replace(" ", "-")
    markdown = row["markdown"]
    if format in {"markdown", "md"}:
        data = markdown.encode("utf-8")
        return StreamingResponse(
            BytesIO(data),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{title}.md"'},
        )
    if format == "docx":
        from docx import Document

        doc = Document()
        for line in markdown.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
        )
    if format == "pdf":
        data = _markdown_to_minimal_pdf(markdown)
        return StreamingResponse(
            BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{title}.pdf"'},
        )
    raise HTTPException(status_code=400, detail="format must be markdown, docx, or pdf")


@app.post("/api/gap/analyze")
async def create_gap_analysis(payload: GapAnalysisRequest) -> dict:
    if payload.resume_id is not None:
        profile_id, profile = _resume_as_profile(payload.resume_id)
    elif payload.profile_id is not None:
        profile_id = payload.profile_id
        profile = _row_json("profiles", payload.profile_id, "profile_json")
    else:
        raise HTTPException(status_code=400, detail="profile_id or resume_id is required")
    role = _row_json("role_analyses", payload.role_analysis_id, "analysis_json")
    analysis = await analyze_gap(profile, role)
    if payload.resume_id is not None:
        analysis["resume_id"] = payload.resume_id
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO gap_analyses (profile_id, role_analysis_id, analysis_json, created_at)
            VALUES (?, ?, ?, ?)
            """,
            (profile_id, payload.role_analysis_id, encode_json(analysis), now_iso()),
        )
    return {"id": int(cur.lastrowid), "analysis": analysis}


@app.get("/api/gaps")
def list_gap_analyses() -> dict:
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT id, profile_id, role_analysis_id, analysis_json, created_at
            FROM gap_analyses ORDER BY created_at DESC
            """
        ).fetchall()
    return {
        "gaps": [
            {
                "id": row["id"],
                "profile_id": row["profile_id"],
                "role_analysis_id": row["role_analysis_id"],
                "analysis": decode_json(row["analysis_json"]),
                "created_at": row["created_at"],
            }
            for row in rows
        ]
    }


@app.get("/api/gap/{gap_id}")
def get_gap_analysis(gap_id: int) -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM gap_analyses WHERE id = ?", (gap_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Gap analysis not found")
    return {
        "id": row["id"],
        "profile_id": row["profile_id"],
        "role_analysis_id": row["role_analysis_id"],
        "analysis": decode_json(row["analysis_json"]),
        "created_at": row["created_at"],
    }


@app.put("/api/gap/{gap_id}")
def update_gap_analysis(gap_id: int, payload: GapAnalysisUpdateRequest) -> dict:
    with get_db() as conn:
        row = conn.execute("SELECT id FROM gap_analyses WHERE id = ?", (gap_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Gap analysis not found")
        conn.execute(
            "UPDATE gap_analyses SET analysis_json = ? WHERE id = ?",
            (encode_json(payload.analysis), gap_id),
        )
    return {"saved": True, "id": gap_id}


@app.delete("/api/gap/{gap_id}")
def delete_gap_analysis(gap_id: int) -> dict:
    with get_db() as conn:
        cur = conn.execute("DELETE FROM gap_analyses WHERE id = ?", (gap_id,))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Gap analysis not found")
    return {"deleted": True, "id": gap_id}


@app.post("/api/improvement/generate")
async def create_improvement_plan(payload: dict) -> dict:
    gap_id = payload.get("gap_analysis_id")
    if not gap_id:
        raise HTTPException(status_code=400, detail="gap_analysis_id is required")
    
    with get_db() as conn:
        gap_row = conn.execute(
            "SELECT profile_id, role_analysis_id, analysis_json FROM gap_analyses WHERE id = ?",
            (gap_id,)
        ).fetchone()
        if not gap_row:
            raise HTTPException(status_code=404, detail="Gap analysis not found")
        
        profile_row = conn.execute(
            "SELECT profile_json FROM profiles WHERE id = ?",
            (gap_row["profile_id"],)
        ).fetchone()
        
        role_row = conn.execute(
            "SELECT analysis_json FROM role_analyses WHERE id = ?",
            (gap_row["role_analysis_id"],)
        ).fetchone()
    
    gap_analysis = decode_json(gap_row["analysis_json"])
    profile = decode_json(profile_row["profile_json"]) if profile_row else {}
    role = decode_json(role_row["analysis_json"]) if role_row else {}
    
    plan = await generate_improvement_plan(gap_analysis, profile, role)
    
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO improvement_plans (gap_analysis_id, profile_id, role_analysis_id, plan_json, created_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (gap_id, gap_row["profile_id"], gap_row["role_analysis_id"], encode_json(plan), now_iso())
        )
        plan_id = int(cur.lastrowid)
    
    return {"id": plan_id, "gap_analysis_id": gap_id, "plan": plan}


@app.get("/api/improvement/plans")
def list_improvement_plans() -> dict:
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT id, gap_analysis_id, profile_id, role_analysis_id, plan_json, created_at
            FROM improvement_plans
            ORDER BY created_at DESC
            """
        ).fetchall()
    
    plans = []
    for row in rows:
        plans.append({
            "id": row["id"],
            "gap_analysis_id": row["gap_analysis_id"],
            "profile_id": row["profile_id"],
            "role_analysis_id": row["role_analysis_id"],
            "plan": decode_json(row["plan_json"]),
            "created_at": row["created_at"],
        })
    
    return {"plans": plans}


@app.delete("/api/improvement/{plan_id}")
def delete_improvement_plan(plan_id: int) -> dict:
    with get_db() as conn:
        cur = conn.execute("DELETE FROM improvement_plans WHERE id = ?", (plan_id,))
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Improvement plan not found")
    return {"deleted": True, "id": plan_id}


@app.post("/api/interview/start")
def start_interview(payload: InterviewStartRequest) -> dict:
    profile_id = payload.profile_id
    context_note = ""
    if payload.resume_id is not None:
        profile_id, resume_profile = _resume_as_profile(payload.resume_id)
        context_note = f" Use resume #{payload.resume_id}: {resume_profile['resume_title']}."
    question = first_question(payload.mode, payload.difficulty) + context_note
    history = [{"role": "assistant", "content": question}]
    timestamp = now_iso()
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO interview_sessions
            (profile_id, role_analysis_id, mode, difficulty, question_count, history_json, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                profile_id,
                payload.role_analysis_id,
                payload.mode,
                payload.difficulty,
                payload.question_count,
                encode_json(history),
                timestamp,
                timestamp,
            ),
        )
    return {
        "session_id": int(cur.lastrowid),
        "question": question,
        "question_count": payload.question_count,
        "question_number": 1,
        "history": history,
    }


@app.post("/api/interview/answer")
async def answer_interview(payload: InterviewAnswerRequest) -> dict:
    with get_db() as conn:
        row = conn.execute(
            "SELECT history_json, question_count, mode, difficulty FROM interview_sessions WHERE id = ?",
            (payload.session_id,),
        ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Interview session not found")

    history = decode_json(row[0])
    question_count = row[1]
    mode = row[2]
    difficulty = row[3]

    # Calculate current question number (each Q&A pair = 2 entries)
    question_number = (len(history) + 1) // 2

    # Handle skip action
    if payload.action == "skip":
        history.append({"role": "user", "content": "[SKIPPED]"})
        question_number = (len(history) + 1) // 2

        if question_number > question_count:
            feedback = await generate_session_feedback(history, question_count)
            with get_db() as conn:
                conn.execute(
                    "UPDATE interview_sessions SET history_json = ?, updated_at = ? WHERE id = ?",
                    (encode_json(history), now_iso(), payload.session_id),
                )
            return {
                "session_complete": True,
                "is_final_question": True,
                "session_feedback": feedback,
            }
        else:
            history.append({"role": "assistant", "content": "Next question..."})
            with get_db() as conn:
                conn.execute(
                    "UPDATE interview_sessions SET history_json = ?, updated_at = ? WHERE id = ?",
                    (encode_json(history), now_iso(), payload.session_id),
                )
            next_question = next((item["content"] for item in reversed(history) if item["role"] == "assistant"), "")
            return {
                "next_question": next_question,
                "question_number": question_number,
                "question_count": question_count,
                "is_final_question": False,
            }

    # Handle exit action
    if payload.action == "exit":
        feedback = await generate_session_feedback(history, question_count)
        with get_db() as conn:
            conn.execute(
                "UPDATE interview_sessions SET history_json = ?, updated_at = ? WHERE id = ?",
                (encode_json(history), now_iso(), payload.session_id),
            )
        return {
            "session_complete": True,
            "is_final_question": True,
            "session_feedback": feedback,
            "exit_reason": "User exited early",
        }

    # Normal submit action
    if not payload.answer.strip():
        raise HTTPException(status_code=400, detail="Answer cannot be empty for submit action")

    history.append({"role": "user", "content": payload.answer})

    last_question = next((item["content"] for item in reversed(history) if item["role"] == "assistant"), "")
    evaluation = await evaluate_answer(mode, difficulty, last_question, payload.answer)
    evaluation_entry = {
        "role": "assistant",
        "content": evaluation["next_question"],
        "evaluation": evaluation,
    }
    history.append(evaluation_entry)

    question_number = (len(history) + 1) // 2
    is_final = question_number > question_count

    with get_db() as conn:
        conn.execute(
            "UPDATE interview_sessions SET history_json = ?, updated_at = ? WHERE id = ?",
            (encode_json(history), now_iso(), payload.session_id),
        )

    if is_final:
        feedback = await generate_session_feedback(history, question_count)
        return {
            "evaluation": evaluation,
            "session_complete": True,
            "is_final_question": True,
            "session_feedback": feedback,
        }
    else:
        return {
            "evaluation": evaluation,
            "next_question": evaluation["next_question"],
            "question_number": question_number,
            "question_count": question_count,
            "is_final_question": False,
        }
